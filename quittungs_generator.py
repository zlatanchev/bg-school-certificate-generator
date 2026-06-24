# -*- coding: utf-8 -*-
"""
Quittungs-Generator für Schulgebühren - Threading, Cancelling & PDF/Word Edition
"""

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
import os
import sys
import shutil
import threading
from datetime import datetime
from num2words import num2words

# --- NEU: FIX FÜR DIE .EXE DATEI (NOCONSOLE) ---
# Verhindert den Absturz, weil docx2pdf versucht, in eine unsichtbare Konsole zu schreiben
class DummyOutput:
    def write(self, x): pass
    def flush(self): pass

if sys.stdout is None:
    sys.stdout = DummyOutput()
if sys.stderr is None:
    sys.stderr = DummyOutput()
# -----------------------------------------------

# --- IMPORTE FÜR PDF ---
try:
    from docx2pdf import convert
    from pypdf import PdfWriter
    import pythoncom  # WICHTIG: Wird für Windows COM-Threading benötigt
    HAS_PDF_TOOLS = True
except ImportError:
    HAS_PDF_TOOLS = False

# Versuche Pillow zu importieren (für die Bildskalierung)
try:
    from PIL import Image, ImageTk
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

# Globale Variable für den Abbruch (Cancelling)
cancel_flag = False

def initialize_paths():
    try:
        if getattr(sys, 'frozen', False):
            base_path = os.path.dirname(sys.executable)
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))

        default_files = {
            "schuelerliste": "schuelerliste.xlsx",
            "preise": "preise.xlsx",
            "template": "Quittung-Template.docx",
            "logo": "logo.png"
        }

        path_schueler = os.path.join(base_path, default_files["schuelerliste"])
        if os.path.exists(path_schueler):
            excel_path_var.set(path_schueler)

        path_preise = os.path.join(base_path, default_files["preise"])
        if os.path.exists(path_preise):
            prices_path_var.set(path_preise)
        
        path_template = os.path.join(base_path, default_files["template"])
        if os.path.exists(path_template):
            template_path_var.set(path_template)
            
        path_logo = os.path.join(base_path, default_files["logo"])
        if os.path.exists(path_logo):
            logo_path_var.set(path_logo)

        out_dir = os.path.join(base_path, "out")
        output_dir_var.set(out_dir)
            
    except Exception as e:
        print(f"Fehler bei der Initialisierung der Pfade: {e}")

def docx_replace_text(doc_obj, old_text, new_text):
    for p in doc_obj.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text
                    for j in range(i + 1, len(inline)):
                        if old_text in inline[j].text:
                            inline[j].text = inline[j].text.replace(old_text, "")
                    break 

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_text(cell, old_text, new_text)

def load_prices(filepath):
    price_sheets = pd.read_excel(filepath, sheet_name=None)
    
    fee_df = price_sheets['Gebuehren']
    fee_df.columns = fee_df.columns.str.strip() 
    child_fees = pd.Series(fee_df.Betrag.values, index=fee_df.Kind_Nr).to_dict()

    contribution_df = price_sheets['Beitraege']
    contribution_df.columns = contribution_df.columns.str.strip()
    membership_fee = contribution_df[contribution_df.Posten == 'Mitgliedsbeitrag']['Betrag'].iloc[0]
    
    config_df = price_sheets['Konfiguration']
    config_df.columns = config_df.columns.str.strip()
    school_year = config_df[config_df.Eigenschaft == 'Schuljahr']['Wert'].iloc[0]
    
    return child_fees, float(membership_fee), str(school_year)

def toggle_buttons(running=False):
    """Aktiviert/Deaktiviert Buttons basierend auf dem Thread-Status"""
    state = tk.DISABLED if running else tk.NORMAL
    cancel_state = tk.NORMAL if running else tk.DISABLED
    
    btn_generate_word.config(state=state)
    btn_generate_pdf.config(state=state)
    btn_cancel.config(state=cancel_state)

def cancel_process():
    """Setzt das Abbruch-Flag"""
    global cancel_flag
    cancel_flag = True
    status_var.set("Abbruch wird eingeleitet... Bitte warten.")
    btn_cancel.config(state=tk.DISABLED)

def show_message_threadsafe(title, msg, is_error=False, is_warning=False):
    if is_error:
        messagebox.showerror(title, msg)
    elif is_warning:
        messagebox.showwarning(title, msg)
    else:
        messagebox.showinfo(title, msg)


# ==========================================
# PHASE 1: WORD-DOKUMENTE GENERIEREN (THREAD)
# ==========================================
def start_word_generation():
    global cancel_flag
    cancel_flag = False
    
    excel_path = excel_path_var.get()
    template_path = template_path_var.get()
    prices_path = prices_path_var.get()
    output_dir = output_dir_var.get()
    
    if not all([excel_path, template_path, prices_path, output_dir]):
        messagebox.showerror("Fehler", "Bitte alle Pfade auswählen!")
        return

    toggle_buttons(running=True)
    progress_var.set(0)
    
    threading.Thread(target=generate_word_receipts_task, args=(excel_path, template_path, prices_path, output_dir), daemon=True).start()

def generate_word_receipts_task(excel_path, template_path, prices_path, output_dir):
    global cancel_flag
    errors_found = []
    class_folders = set()
    quittungs_nr = 1 # Der globale Zähler
    
    try:
        child_fees, membership_fee, school_year = load_prices(prices_path)
        
        df = pd.read_excel(excel_path)
        df.dropna(subset=['Eltern 1 - Emailadresse', 'Name Kind'], inplace=True)
        df['Eltern 1 - Emailadresse'] = df['Eltern 1 - Emailadresse'].astype(str).str.strip()
        
        # DataFrame nach Klassen sortieren, damit die Nummern kontinuierlich steigen
        df['In Klasse Sortierung'] = df['In Klasse'].astype(str)
        df.sort_values(by='In Klasse Sortierung', inplace=True)
        
        grouped = df.groupby('Eltern 1 - Emailadresse', sort=False)
        
        total_parents = len(grouped)
        
        root.after(0, lambda: progress_bar.config(maximum=total_parents))
        root.after(0, lambda: status_var.set(f"Starte Generierung von {total_parents} Word-Quittungen..."))
        
        current_progress = 0

        for parent_email, group in grouped:
            if cancel_flag:
                break

            try:
                parent_full_name = str(group['Eltern 1 - Name'].iloc[0]).strip()
                
                is_group_valid = True
                for index, row in group.iterrows():
                    kind_value = row['Name Kind']
                    if not isinstance(kind_value, str):
                        excel_row_number = index + 2
                        errors_found.append(f"Mitglied: '{parent_full_name}' ({parent_email})\nGrund: Ungültiger Datentyp in Spalte 'Name Kind' (Zeile {excel_row_number}).")
                        is_group_valid = False
                        break 
                
                if not is_group_valid: continue 
                if group['In Klasse'].isin(['Warteliste', '', ' ']).any(): continue

                doc = Document(template_path)
                num_children = len(group)
                
                kinder_liste = [str(name) for name in group['Name Kind']]
                if len(kinder_liste) > 2:
                    children_names = ", ".join(kinder_liste[:-1]) + " und " + kinder_liste[-1]
                else:
                    children_names = " und ".join(kinder_liste)

                total_school_fee = sum(child_fees.get(i, 0) for i in range(1, num_children + 1))
                total_amount = total_school_fee + membership_fee
                
                klasse = str(group['In Klasse'].iloc[0])
                safe_klasse = klasse.replace("/", "_").replace("\\", "_")
                
                # GLOBAL & REIN NUMERISCH
                eindeutige_nummer = f"{quittungs_nr:03d}"
                
                replacements = {
                    "{{ELTERN_NAME}}": parent_full_name,
                    "{{KINDER_NAMEN}}": children_names, 
                    "{{NR}}": eindeutige_nummer, 
                    "{{DATUM}}": datetime.now().strftime("%d.%m.%Y"), 
                    "{{SCHULJAHR}}": str(school_year),
                    "{{BETRAG_GEBUEHR}}": f"{total_school_fee:,.2f} EUR".replace(",", "X").replace(".", ",").replace("X", "."),
                    "{{GESAMTBETRAG}}": f"{total_amount:,.2f} EUR".replace(",", "X").replace(".", ",").replace("X", "."),
                    "{{BETRAG_GEBUEHR_WORT}}": f"{num2words(int(total_school_fee), lang='de')} Euro", 
                    "{{GESAMTBETRAG_WORT}}": f"{num2words(int(total_amount), lang='de')} Euro",
                    "{{BETRAG_MITGLIED}}": f"{membership_fee:,.2f} EUR".replace(",", "X").replace(".", ",").replace("X", "."),
                    "{{BETRAG_MITGLIED_WORT}}": f"{num2words(int(membership_fee), lang='de')} Euro",
                }

                for old, new in replacements.items():
                    docx_replace_text(doc, old, str(new))

                outdir_class = os.path.join(output_dir, safe_klasse)
                
                if not os.path.exists(outdir_class):
                    os.makedirs(outdir_class)
                class_folders.add(outdir_class)

                safe_parent_name = parent_full_name.replace(" ", "_").replace("/", "_")
                
                # --- NEU: Dateiname mit Erstellungsjahr, Quittungs-Nummer und Name ---
                jahr_erstellung = datetime.now().strftime("%Y")
                dateiname = f"{jahr_erstellung}_Quittung_{eindeutige_nummer}_{safe_parent_name}.docx"
                output_filename = os.path.join(outdir_class, dateiname)
                
                doc.save(output_filename)

                # Zähler für die nächste Familie erhöhen
                quittungs_nr += 1
                
                current_progress += 1
                root.after(0, progress_var.set, current_progress)
                root.after(0, status_var.set, f"Erstelle Word-Dokumente... ({current_progress}/{total_parents})")

            except Exception as e:
                errors_found.append(f"Mitglied: '{parent_email}'\nGrund: Unerwarteter Fehler -> {e}")
                continue 

        if cancel_flag:
            root.after(0, status_var.set, "Prozess durch Benutzer abgebrochen.")
            return

        # Abschlussmeldung Word
        root.after(0, progress_var.set, total_parents)
        
        generierte_quittungen = quittungs_nr - 1
        anzahl_klassen = len(class_folders)

        zusammenfassung = (
            f"Statistik:\n"
            f"➜ {generierte_quittungen} Quittungen (Familien) erstellt.\n"
            f"➜ Aufgeteilt in {anzahl_klassen} verschiedene Klassen."
        )
        
        if not errors_found:
            root.after(0, status_var.set, "Word-Generierung erfolgreich abgeschlossen!")
            root.after(0, show_message_threadsafe, "Schritt 1 abgeschlossen", f"Word-Dateien erfolgreich generiert!\n\n{zusammenfassung}\n\nDu kannst die Dateien nun im Ausgabeordner kontrollieren und bei Bedarf anpassen, bevor du Schritt 2 ausführst.")
        else:
            root.after(0, status_var.set, "Mit Warnungen abgeschlossen.")
            error_summary = "\n\n------------------------------------\n\n".join(errors_found)
            final_message = (
                f"Word-Dateien wurden generiert.\n\n{zusammenfassung}\n\n"
                f"Es gab jedoch Probleme/Fehler:\n\n{error_summary}"
            )
            root.after(0, show_message_threadsafe, "Word-Generierung (mit Warnungen)", final_message, False, True)

    except Exception as e:
        root.after(0, status_var.set, "Kritischer Fehler aufgetreten!")
        root.after(0, show_message_threadsafe, "Kritischer Fehler", f"Ein Fehler hat die Verarbeitung gestoppt:\n{e}", True)
    finally:
        root.after(0, toggle_buttons, False)


# ==========================================
# PHASE 2: PDFS GENERIEREN (THREAD)
# ==========================================
def start_pdf_generation():
    global cancel_flag
    
    if not HAS_PDF_TOOLS:
        messagebox.showerror("Fehlende Pakete", "Bitte installiere die PDF-Erweiterungen im Terminal:\n\npip install docx2pdf pypdf pythoncom")
        return

    output_dir = output_dir_var.get()
    if not output_dir:
        messagebox.showerror("Fehler", "Bitte den Ausgabeordner auswählen!")
        return

    cancel_flag = False
    toggle_buttons(running=True)
    progress_var.set(0)
    
    threading.Thread(target=generate_pdf_receipts_task, args=(output_dir,), daemon=True).start()

def generate_pdf_receipts_task(output_dir):
    global cancel_flag
    errors_found = []

    if HAS_PDF_TOOLS:
        pythoncom.CoInitialize()

    try:
        class_folders = []
        total_docx_files = 0
        
        if os.path.exists(output_dir):
            for element in os.listdir(output_dir):
                element_path = os.path.join(output_dir, element)
                if os.path.isdir(element_path):
                    docx_files = [f for f in os.listdir(element_path) if f.endswith('.docx') and not f.startswith('~')]
                    if docx_files:
                        class_folders.append(element_path)
                        total_docx_files += len(docx_files)

        if not class_folders:
            root.after(0, status_var.set, "Warte auf Start...")
            root.after(0, show_message_threadsafe, "Info", "Keine Klassen-Ordner mit Word-Dateien im Ausgabeordner gefunden.\nBitte führe zuerst Schritt 1 aus.")
            return

        anzahl_klassen = len(class_folders)
        
        root.after(0, lambda: progress_bar.config(maximum=anzahl_klassen))
        root.after(0, lambda: status_var.set(f"Starte PDF-Konvertierung für {anzahl_klassen} Klassen..."))

        current_progress = 0
        pdf_count = 0

        for class_folder in class_folders:
            if cancel_flag:
                break

            klasse_name = os.path.basename(class_folder)
            root.after(0, status_var.set, f"Konvertiere Klasse {klasse_name} in PDFs... ({current_progress + 1}/{anzahl_klassen})")
            
            try:
                # 1. Wandelt den gesamten Ordner in PDFs um
                convert(class_folder)
                
                # 2. PDFs zusammenfügen
                merger = PdfWriter()
                pdf_files = sorted([f for f in os.listdir(class_folder) if f.endswith('.pdf')])
                
                if pdf_files:
                    final_pdf_path = os.path.join(output_dir, f"Sammel_PDF_Klasse_{klasse_name}.pdf")
                    for pdf in pdf_files:
                        merger.append(os.path.join(class_folder, pdf))
                        
                    merger.write(final_pdf_path)
                    merger.close()
                    pdf_count += 1
                
                current_progress += 1
                root.after(0, progress_var.set, current_progress)
                    
            except Exception as e:
                errors_found.append(f"Fehler bei PDF-Erstellung für Klasse {klasse_name}: {e}\n(Ist Microsoft Word geschlossen und bereit?)")

        if cancel_flag:
            root.after(0, status_var.set, "Prozess durch Benutzer abgebrochen.")
            return

        # Abschlussmeldung PDF
        root.after(0, progress_var.set, anzahl_klassen)
        
        zusammenfassung = (
            f"Statistik:\n"
            f"➜ {total_docx_files} einzelne Quittungen verarbeitet.\n"
            f"➜ {total_docx_files} Einzel-PDFs in den Klassenordnern generiert.\n"
            f"➜ {pdf_count} Sammel-PDFs (Klassen) erfolgreich im Ausgabeordner erstellt."
        )

        if not errors_found:
            root.after(0, status_var.set, "PDF-Sammelquittungen erfolgreich generiert!")
            root.after(0, show_message_threadsafe, "Schritt 2 abgeschlossen", f"PDF-Prozess erfolgreich beendet!\n\n{zusammenfassung}")
        else:
            root.after(0, status_var.set, "Mit Warnungen abgeschlossen.")
            error_summary = "\n\n------------------------------------\n\n".join(errors_found)
            final_message = (
                f"PDF-Sammelquittungen wurden generiert.\n\n{zusammenfassung}\n\n"
                f"Es gab jedoch Probleme/Fehler:\n\n{error_summary}"
            )
            root.after(0, show_message_threadsafe, "PDF-Generierung (mit Warnungen)", final_message, False, True)

    except Exception as e:
        root.after(0, status_var.set, "Kritischer Fehler aufgetreten!")
        root.after(0, show_message_threadsafe, "Kritischer Fehler", f"Ein Fehler hat die PDF-Verarbeitung gestoppt:\n{e}", True)
    finally:
        if HAS_PDF_TOOLS:
            pythoncom.CoUninitialize()
        root.after(0, toggle_buttons, False)


# --- GUI Code ---
def select_excel_file():
    filepath = filedialog.askopenfilename(filetypes=[("Excel-Dateien", "*.xlsx *.xls")])
    if filepath:
        excel_path_var.set(filepath)

def select_template_file():
    filepath = filedialog.askopenfilename(filetypes=[("Word-Dokumente", "*.docx")])
    if filepath:
        template_path_var.set(filepath)

def select_prices_file():
    filepath = filedialog.askopenfilename(filetypes=[("Excel-Dateien", "*.xlsx *.xls")])
    if filepath:
        prices_path_var.set(filepath)

def select_output_dir():
    dirpath = filedialog.askdirectory()
    if dirpath:
        output_dir_var.set(dirpath)

root = tk.Tk()
root.title("Quittungs-Generator (PDF & Word Edition - Multi-Thread)")
root.geometry("650x550") 

excel_path_var = tk.StringVar()
template_path_var = tk.StringVar()
prices_path_var = tk.StringVar()
output_dir_var = tk.StringVar()
logo_path_var = tk.StringVar()

frame = tk.Frame(root, padx=10, pady=10)
frame.pack(expand=True, fill=tk.X) 

frame.grid_columnconfigure(0, weight=1)
frame.grid_columnconfigure(1, weight=0)

initialize_paths()

logo_path = logo_path_var.get()
if os.path.exists(logo_path):
    if HAS_PILLOW:
        try:
            img = Image.open(logo_path)
            max_width = 500
            max_height = 120
            try:
                resample_filter = Image.Resampling.LANCZOS 
            except AttributeError:
                resample_filter = Image.ANTIALIAS 

            img.thumbnail((max_width, max_height), resample_filter)
            logo_img = ImageTk.PhotoImage(img)
            root.logo_img = logo_img 
            tk.Label(frame, image=logo_img).grid(row=0, column=0, columnspan=2, pady=(0, 15))
        except Exception as e:
            tk.Label(frame, text="[Fehler bei der Logo-Verarbeitung]").grid(row=0, column=0, columnspan=2, pady=(0, 15))
    else:
        tk.Label(frame, text="[Bitte 'Pillow' installieren für Logo-Skalierung]", fg="red").grid(row=0, column=0, columnspan=2, pady=(0, 15))

# Eingabefelder und Buttons
tk.Label(frame, text="1. Excel-Datei (Schülerliste) auswählen:").grid(row=1, column=0, sticky="w", pady=2)
tk.Entry(frame, textvariable=excel_path_var, width=60).grid(row=2, column=0, padx=(0, 5), sticky="ew")
tk.Button(frame, text="Durchsuchen...", command=select_excel_file).grid(row=2, column=1)

tk.Label(frame, text="2. Excel-Datei (Preise) auswählen:").grid(row=3, column=0, sticky="w", pady=(10, 2))
tk.Entry(frame, textvariable=prices_path_var, width=60).grid(row=4, column=0, padx=(0, 5), sticky="ew")
tk.Button(frame, text="Durchsuchen...", command=select_prices_file).grid(row=4, column=1)

tk.Label(frame, text="3. Word-Vorlagendatei auswählen:").grid(row=5, column=0, sticky="w", pady=(10, 2))
tk.Entry(frame, textvariable=template_path_var, width=60).grid(row=6, column=0, padx=(0, 5), sticky="ew")
tk.Button(frame, text="Durchsuchen...", command=select_template_file).grid(row=6, column=1)

tk.Label(frame, text="4. Ausgabeordner auswählen:").grid(row=7, column=0, sticky="w", pady=(10, 2))
tk.Entry(frame, textvariable=output_dir_var, width=60).grid(row=8, column=0, padx=(0, 5), sticky="ew")
tk.Button(frame, text="Durchsuchen...", command=select_output_dir).grid(row=8, column=1)

# Frame für die Steuerungsknöpfe
button_frame = tk.Frame(frame)
button_frame.grid(row=9, column=0, columnspan=2, pady=(20, 5))

btn_generate_word = tk.Button(button_frame, text="📝 1. Word generieren", font=("Helvetica", 11, "bold"), command=start_word_generation, bg="#2196F3", fg="white")
btn_generate_word.pack(side=tk.LEFT, padx=5, ipadx=5, ipady=5)

btn_generate_pdf = tk.Button(button_frame, text="📄 2. PDFs generieren", font=("Helvetica", 11, "bold"), command=start_pdf_generation, bg="#4CAF50", fg="white")
btn_generate_pdf.pack(side=tk.LEFT, padx=5, ipadx=5, ipady=5)

btn_cancel = tk.Button(button_frame, text="🛑 Abbrechen", font=("Helvetica", 11, "bold"), command=cancel_process, bg="#cc3025", fg="white", state=tk.DISABLED)
btn_cancel.pack(side=tk.LEFT, padx=5, ipadx=5, ipady=5)

# Status-Text
status_var = tk.StringVar()
status_var.set("Warte auf Start...")
tk.Label(frame, textvariable=status_var, fg="blue", font=("Helvetica", 10)).grid(row=10, column=0, columnspan=2, pady=(0, 5))

# Fortschrittsbalken
progress_var = tk.IntVar()
progress_bar = ttk.Progressbar(frame, variable=progress_var, mode='determinate')
progress_bar.grid(row=11, column=0, columnspan=2, sticky="ew", pady=(0, 15))

# Info-Feld
tk.Label(frame, text="Version 25.06.2026; I. Zlat.", font=("Helvetica", 8), fg="gray").grid(row=12, column=0, columnspan=2, pady=(0, 5))

root.mainloop()